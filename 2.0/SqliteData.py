import sqlite3 as lite
import sys
from datetime import datetime
import matplotlib
import matplotlib.pyplot as plt
from os import path
from PIL import Image
import numpy as np
import jieba
#from wordcloud import WordCloud, STOPWORDS
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.shared import RGBColor

#freq 1=day 2=week 3=month
def getTagDate(freq):
    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()

        #user want yesterday's data
        if(freq==1):
            cur.execute("select date('now')");
            for row in cur:
                return row[0]
        
        #user want last week's data
        elif(freq==2):
            if(datetime.today().weekday()==0):
                cur.execute("select date('now', 'weekday 1')");
                for row in cur:
                    return row[0]
            else:
                cur.execute("select date('now', 'weekday 1', '-7 days')");
                for row in cur:
                    return row[0]

        #user want last month's data
        elif(freq==3):
            cur.execute("select date('now', 'start of month')");
            for row in cur:
                return row[0]

def getStartDate(freq):
    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()

        #user want yesterday's data
        if(freq==1):
            cur.execute("select date('now','-1 days')");
            for row in cur:
                return row[0]
        
        #user want last week's data
        elif(freq==2):
            if(datetime.today().weekday()==0):
                cur.execute("select date('now', 'weekday 1','-1 days')");
                for row in cur:
                    return row[0]
            else:
                cur.execute("select date('now', 'weekday 1', '-7 days')");
                for row in cur:
                    return row[0]

        #user want last month's data
        elif(freq==3):
            cur.execute("select date('now', 'start of month','-1 days')");
            for row in cur:
                return row[0]

def getEndDate(freq):
    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()
        if(freq==1):
            cur.execute("select date('now','-1 day')");
            for row in cur:
                return row[0]
        elif(freq==2):
            if(datetime.today().weekday()==0):
                cur.execute("select date('now', 'weekday 1', '-7 days')");
                for row in cur:
                    return row[0]
            else:
                cur.execute("select date('now', 'weekday 1', '-14 days')");
                for row in cur:
                    return row[0]
        elif(freq==3):
            cur.execute("select date('now', 'start of month', '-1 month')");
            for row in cur:
                return row[0]

# index data for index.html ,catch all necessary data 
def index_data(freq):
    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()
        ary=[]
        igotDate=getTagDate(freq)
        
        #default is to get every day's data
        #get tags that user choose
        cur.execute("SELECT words,wordsCount,frequencyNum FROM wordsCount WHERE frequencyNum ={freq} AND daytime = '{getDate}' ORDER BY wordsCount DESC LIMIT 50" .\
        format(freq=freq,getDate=igotDate));
        
        #this for test
        # cur.execute("SELECT words,wordsCount,frequencyNum FROM wordsCount WHERE frequencyNum ={freq} ORDER BY wordsCount DESC LIMIT 50" .\
        #         format(freq=freq));
        
        for row in cur:
            dict={'words':row[0],'wordsCount': row[1],'frequencyNum':row[2]}
            ary.append(dict)
        return ary

#wordCloud data,js need to use it , get wordsCount words wordsCount
def wdCloud_data(start,end):
    ary=[]
    summation = 0
    counter = 0
    Start_date = datetime.strptime(start,'%Y-%m-%d')
    End_date = datetime.strptime(end,'%Y-%m-%d')
    connect = lite.connect('../FBWEB.sqlite')
    cur = connect.cursor()
    a = cur.execute("SELECT summary FROM WebCrawler WHERE dt BETWEEN ? AND ?",(Start_date,End_date,)).fetchall()
    for summ in a:
        ary.append(summ)
    
    stopWords = []
    words_ary = []
    re = {}

    strsummary = str(ary)
    with open('stopWords.txt', 'r', encoding='UTF-8') as file:
        for data in file.readlines():
            data = data.strip()
            stopWords.append(data)

    for word in jieba.cut(strsummary):
        words_ary.append(word)
    from collections import Counter
    words_ary = [word for word in jieba.cut(strsummary)]
    cc = Counter(words_ary)
    hik ={}
    for k, v in cc.most_common(100):
        if k in stopWords:
            continue
        else:
            if len(k) >= 2 and v >= 5:##k = 我、的(少於一個字的詞) V = 詞的最低數量
                hik[k]=hik.get(k,v)
    return hik

#for tag.html  when user click any bottom , show courrect data
def tag_data(words,freq):
    select_ary=[]

    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()
        endDate=getStartDate(freq)
        startDate=getEndDate(freq)
        #select data to dict and the dict will return to tag.html

        cur.execute("SELECT dt,link,title,summary,countMsg FROM webCrawler WHERE date(dt) BETWEEN '{startDate}' AND '{endDate}' AND summary LIKE '%%{tag}%%' ORDER BY countMsg DESC ".\
            format(startDate=startDate,endDate=endDate,tag=words));
        # this for test 
        # cur.execute("SELECT dt,link,title,summary FROM {tt} WHERE summary LIKE '%%{tag}%%' ".\
        #     format(tt=tt,tag=words));

        for row in cur:
            dict={'dt':row[0],"link": row[1],"title": row[2],'summary':row[3],'countMsg':row[4]}
            select_ary.append(dict)
        return select_ary                   

def gChart_data(freq):
    with lite.connect('../FBWEB.sqlite') as conn:
        cur = conn.cursor()
        ary=[]
        igotDate=getTagDate(freq)
        
        #default is to get every day's data
        #get tags that user choose
        cur.execute("SELECT words,wordsCount,frequencyNum FROM wordsCount WHERE frequencyNum ={freq} AND daytime = '{getDate}' ORDER BY wordsCount DESC LIMIT 10".format(freq=freq,getDate=igotDate));

        # this for test
        # cur.execute("SELECT words,wordsCount FROM wordsCount WHERE frequencyNum ={freq} ORDER BY wordsCount DESC LIMIT 10".\
        #         format(freq=freq));
        
        for row in cur:
            dict={'words':row[0],'wordsCount': row[1]}
            ary.append(dict)
        return ary          

def keywordsearch(keywordstring):
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    d = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    total = []
    Id = []
    keywords = [x.strip() for x in keywordstring.split(',')]
    for x in keywords:
        t = c.execute("SELECT Id FROM News WHERE title LIKE '%{}%' OR summary LIKE '%{}%'".format(x,x)).fetchall()
        for tt in t:
            if tt not in Id:
                Id.append(tt)          
    for i in Id:
        d.execute("SELECT dt,title,link,summary FROM News WHERE Id = ?",(i,))
        for row in d:
            dict={'dt':row[0],"title": row[1],"link": row[2],'summary':row[3]}
            total.append(dict)
    return total,keywords

def normalnews():
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    d = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    total = []
    today = datetime.today()
    today = datetime.strftime(today,'%Y/%m/%d')
    d.execute("SELECT dt,title,link,summary FROM News WHERE dt = ?",(today,))
    for row in d:
        dict={'dt':row[0],"title": row[1],"link": row[2],'summary':row[3]}
        total.append(dict)
    return total

def gettitle(links):
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    ti = c.execute("SELECT title FROM News WHERE link = ?",(links,)).fetchone()
    return ti

def getsource(links):
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    source = c.execute("SELECT website FROM News WHERE link = ?",(links,)).fetchone()
    return source

def getcontent(links):
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    content = c.execute("SELECT content FROM News WHERE link = ?",(links,)).fetchone()
    return content

def getdate(links):
    con = lite.connect('../FBWEB.sqlite')
    c = con.cursor()
    c.row_factory = lambda cursor, row: row[0]
    date = c.execute("SELECT dt FROM News WHERE link = ?",(links,)).fetchone()
    source = c.execute("SELECT website FROM News WHERE link = ?",(links,)).fetchone()
    text = date + ' ' + source
    return text

def spliter(lis):
    li = []
    for item in lis:
        li.append(item)
        if 'http' not in item:
            li = li[::-1]
            return(li)
        
def createtable(change):
    doc = Document('demo.docx')
    today = datetime.today()
    today = datetime.strftime(today,'%Y/%m/%d')
    doc.add_paragraph(today)
    for lis in change:
        table = doc.add_table(rows = len(lis), cols = 2,style = 'Light Grid Accent 5')

        row = table.rows[0]
        a, b = row.cells[:2]
        a.merge(b)

        shading_elm = parse_xml(r'<w:shd {} w:fill="003377"/>'.format(nsdecls('w')))
        row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        t = row.cells[0]

        for i in range(1,len(lis)):
            row = table.rows[i]
            row.cells[0].text = getsource(lis[i])
            row.cells[1].text = gettitle(lis[i]).strip()
            row.cells[1].hyperlink = lis[i]
            print(lis[i])


        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(15)

        for i in range(1,len(lis)):
            row = table.rows[i]
            p = row.cells[1].paragraphs
            row.cells[1].width = Inches(15)
            for pp in p:
                for run in pp.runs:
                    run.underline = True
                    run.font.color.rgb = RGBColor(35, 79, 203)

        t.text = lis[0]                
        f = t.paragraphs[0]
        k = f.runs
        font = k[0].font
        font.size= Pt(20)
    
    doc.add_paragraph()
    doc.add_paragraph()
#內文表格
    for lis in change:
        table = doc.add_table(rows = 1, cols = 1,style = 'Table Grid')
        row = table.rows[0]
        shading_elm = parse_xml(r'<w:shd {} w:fill="003377"/>'.format(nsdecls('w')))
        row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        t = row.cells[0]
        t.text = lis[0]                
        f = t.paragraphs[0]
        k = f.runs
        font = k[0].font
        font.size= Pt(18)
        for i in range(1,len(lis)):
            table = doc.add_table(rows = 3, cols = 1,style = 'Table Grid')
            #row = table.rows[0]
            table.rows[0].cells[0].text = gettitle(lis[i]).strip()
            table.rows[1].cells[0].text = getdate(lis[i])
            table.rows[2].cells[0].text = getcontent(lis[i])
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(12)
                            
            t = table.rows[0].cells[0]
            f = t.paragraphs[0]
            k = f.runs
            k[0].bold = True
            font = k[0].font
            font.size= Pt(15)
        doc.add_paragraph()
        
    doc.save('富邦人壽新聞重點摘錄.docx')
    

def generateword(test):
    newtest = test[::-1]
    change = []
    counter = 0
    for i in range(0,len(test)):
        if 'http' not in test[i]:
            counter = counter + 1
    new = spliter(newtest)

    for i in range(0,counter):
        new = spliter(newtest)
        test1 = []
        for item in newtest:
            if item not in new:
                test1.append(item)
        newtest = test1
        change.append(new)

    change = change[::-1]
    createtable(change)